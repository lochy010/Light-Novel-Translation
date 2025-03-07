# file_processor\file_handler.py
"""
文件处理模块
功能：实现文档的智能分块、格式解析和输出保存
核心机制：
- 支持DOCX/TXT格式的解析与生成
- 动态分块算法适配多语言特性
- 自动检测文件编码格式
- 保留文档结构并生成符合排版规范的输出文件
"""

import os
import re
import hashlib
import logging
import uuid
import chardet
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 导入配置文件
from config.settings import FILE_HANDLER_CONFIG

# 初始化模块级日志记录器
logger = logging.getLogger("FileHandler")


def get_file_hash(file_path):
    """
    生成文件唯一特征码（MD5哈希值）

    实现原理：
    - 读取文件二进制内容计算MD5摘要
    - 用于会话缓存管理和文件版本标识

    :param file_path: 文件绝对路径
    :return: 32位十六进制哈希字符串
    :raises IOError: 文件读取失败时抛出
    """
    logger.debug(f"[特征码生成] 开始处理文件: {file_path}")
    try:
        with open(file_path, 'rb') as f:
            # 计算文件内容的MD5哈希值
            file_hash = hashlib.md5(f.read()).hexdigest()
            logger.info(f"[特征码生成] 成功生成 | 文件: {os.path.basename(file_path)} | 哈希: {file_hash[:8]}...")
            return file_hash
    except Exception as e:
        logger.error(f"[特征码生成] 处理失败: {str(e)}", exc_info=True)
        raise


def dynamic_split(text, target_lang, max_tokens=FILE_HANDLER_CONFIG["CHUNKING"]["DEFAULT_MAX_TOKENS"]):
    """
    增强型动态分块算法

    核心特性：
    - 语言自适应的分句规则
    - 连接词感知的分块保护
    - 亚洲语言缓冲系数调节
    - 上下文连贯性保持

    :param text: 原始文本内容（段落列表）
    :param target_lang: 目标语言代码（zh/en/ja/ko）
    :param max_tokens: 单块最大字节数（默认从配置文件中读取）
    :return: 分块后的文本列表
    """
    logger.info(f"[智能分块] 启动分块处理 | 目标语言: {target_lang} | 最大长度: {max_tokens}字节")

    # 新增：统一输入为段落列表
    if isinstance(text, str):
        paragraphs = [text]
    else:
        paragraphs = text

    # 从配置文件中获取语言特定配置
    lang_config = FILE_HANDLER_CONFIG["CHUNKING"]["LANG_CONFIG"]
    # 获取当前语言配置，默认使用英语配置
    config = lang_config.get(target_lang, lang_config['en'])
    logger.debug(f"[智能分块] 应用分块配置: {config}")

    chunks = []
    current_chunk = []
    current_length = 0
    # 从配置文件中获取亚洲语言缓冲系数
    factor = FILE_HANDLER_CONFIG["CHUNKING"]["ASIAN_BUFFER_FACTOR"] if target_lang in ['ja', 'zh'] else 1.0

    # 段落级分块处理
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # 计算当前段落字节长度
        para_length = len(para.encode('utf-8'))

        # 段落超过最大长度时强制分割
        if para_length > max_tokens * factor:
            sub_paras = _split_oversized_para(para, target_lang, max_tokens, config)
            for sub in sub_paras:
                sub_length = len(sub.encode('utf-8'))
                if current_length + sub_length > max_tokens * factor:
                    if current_chunk:
                        chunks.append(''.join(current_chunk))
                        current_chunk = []
                        current_length = 0
                current_chunk.append(sub)
                current_length += sub_length
        else:
            if current_length + para_length > max_tokens * factor:
                if current_chunk:
                    chunks.append(''.join(current_chunk))
                    current_chunk = []
                    current_length = 0
            current_chunk.append(para)
            current_length += para_length

    # 处理剩余内容
    if current_chunk:
        chunks.append(''.join(current_chunk))

    logger.info(f"[智能分块] 处理完成 | 原始段落数: {len(paragraphs)} → 分块数: {len(chunks)}")
    return chunks


def _split_oversized_para(para, target_lang, max_tokens, config):
    """处理超大段落的分割（句子级降级）"""
    # 构建正则表达式模式
    sentence_end = config['sentence_end'] + r'\s*'
    connectors = config['connectors']

    # 处理连接词后的错误分割问题
    for conn in connectors:
        para = re.sub(rf'({conn}[。．!?])', rf'\1##NO_SPLIT_{uuid.uuid4().hex}##', para)

    # 执行初步分句
    sentences = re.split(sentence_end, para)

    # 还原临时标记为原始内容
    sentences = [s.replace(f'##NO_SPLIT_{uuid.uuid4().hex}##', '') for s in sentences]

    # 过滤空句子并保留结尾标点
    return [s + '\n' for s in sentences if s.strip()]


def extract_text_from_docx(docx_path):
    """
    解析DOCX文档内容

    实现特性：
    - 保留段落结构
    - 过滤空段落
    - 维护原始换行符

    :param docx_path: DOCX文件路径
    :return: 段落列表
    :raises ValueError: 文件解析失败时抛出
    """
    logger.info(f"[DOCX解析] 开始处理文件: {docx_path}")
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                # 保留段落换行符
                full_text.append(para.text + '\n')
        logger.debug(f"[DOCX解析] 成功解析 | 段落数: {len(full_text)}")
        return full_text
    except Exception as e:
        logger.error(f"[DOCX解析] 处理失败: {str(e)}", exc_info=True)
        raise


def extract_text_from_txt(txt_path):
    """
    改进版TXT文件读取函数，确保常见语言（英文、中文、日文等）均能正常读入。

    修改说明：
    1. 以二进制模式读取文件，避免使用系统默认编码。
    2. 检测文件的BOM标记，如果存在则直接使用对应的UTF系列编码。
    3. 没有BOM时利用chardet检测编码。
    4. 构造编码优先级列表，包括常见的'utf-8-sig'、'utf-8'、'cp932'、'euc-jp'、'shift_jis'、'iso-2022-jp'、'gb18030'、'latin-1'等，并加入检测到的编码。
    5. 按优先级逐个尝试解码，直到成功解码为止。
    6. 最后按行分割文本，保留原始换行符便于后续处理。
    """
    logger.info(f"[TXT解析] 开始处理文件: {txt_path}")

    try:
        # 以二进制模式读取文件，避免依赖系统默认编码
        with open(txt_path, 'rb') as f:
            raw_data = f.read()
        if not raw_data:
            logger.warning("[TXT解析-改进] 文件为空")
            return []

        # 优先检测BOM标记
        bom_encodings = [
            (b'\xff\xfe\x00\x00', 'utf-32-le'),
            (b'\x00\x00\xfe\xff', 'utf-32-be'),
            (b'\xff\xfe', 'utf-16-le'),
            (b'\xfe\xff', 'utf-16-be'),
            (b'\xef\xbb\xbf', 'utf-8-sig'),
        ]
        detected_encoding = None
        for bom, encoding in bom_encodings:
            if raw_data.startswith(bom):
                detected_encoding = encoding
                logger.info(f"[TXT解析] 检测到BOM，编码为: {encoding}")
                break

        # 若无BOM，则使用chardet检测文件编码
        if not detected_encoding:
            detection = chardet.detect(raw_data)
            detected_encoding = detection['encoding']
            logger.info(f"[TXT解析] chardet检测编码结果: {detected_encoding}")

        # 构造编码优先级列表，覆盖常见语言
        encoding_list = [
            'utf-8-sig',  # 带BOM的UTF-8优先
            'utf-8',  # 标准UTF-8
            'cp932',  # 日语常见编码
            'euc-jp',  # 日语扩展编码
            detected_encoding,
            'shift_jis',
            'iso-2022-jp',
            'gb18030',  # 中文编码后备
            'latin-1'  # 最后保障
        ]
        # 去重并排除无效的编码
        encodings = []
        seen = set()
        for enc in encoding_list:
            if enc and enc.lower() != 'ascii' and enc not in seen:
                encodings.append(enc)
                seen.add(enc)

        # 逐个尝试解码
        text = None
        used_encoding = None
        for enc in encodings:
            try:
                text = raw_data.decode(enc, errors='strict')
                used_encoding = enc
                logger.info(f"[TXT解析] 成功使用编码: {enc}")
                break
            except (UnicodeDecodeError, LookupError):
                logger.debug(f"[TXT解析] 编码尝试失败: {enc}")
                continue

        if text is None:
            raise UnicodeDecodeError("无法使用指定编码解码文件。")

        # 按行分割文本，并保留原始换行符
        paragraphs = text.splitlines(keepends=True)
        logger.debug(f"[TXT解析-改进] 分割后段落数: {len(paragraphs)}")
        return paragraphs

    except Exception as e:
        logger.error(f"[TXT解析-改进] 处理失败: {str(e)}", exc_info=True)
        raise


def save_as_word(content, output_path):
    """
    生成符合排版规范的Word文档

    排版规范：
    - 首行缩进24磅
    - 两端对齐
    - 西文Times New Roman字体
    - 东亚字符楷体
    - 固定行距1倍

    :param content: 内容列表（按分块顺序）
    :param output_path: 输出文件路径
    :raises PermissionError: 文件写入权限不足时抛出
    """
    logger.info(f"[Word生成] 开始保存文件: {output_path}")
    try:
        doc = Document()
        for chunk in content:
            # 按原始段落结构写入
            for para_text in chunk.split('\n'):
                if para_text.strip():
                    p = doc.add_paragraph()
                    # 设置段落格式
                    p_format = p.paragraph_format
                    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
                    p_format.line_spacing = FILE_HANDLER_CONFIG["DOCX_STYLE"]["LINE_SPACING"]  # 固定行距
                    p_format.first_line_indent = FILE_HANDLER_CONFIG["DOCX_STYLE"]["INDENT"]  # 首行缩进

                    # 设置字体样式
                    run = p.add_run(para_text)
                    run.font.name = FILE_HANDLER_CONFIG["DOCX_STYLE"]["FONTS"]["western"]  # 西文字体
                    run.font.size = FILE_HANDLER_CONFIG["DOCX_STYLE"]["FONT_SIZE"]
                    # 东亚字符字体设置
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),
                                                FILE_HANDLER_CONFIG["DOCX_STYLE"]["FONTS"]["east_asian"])

        doc.save(output_path)
        logger.debug(f"[Word生成] 保存成功 | 文件大小: {os.path.getsize(output_path) // 1024}KB")
    except Exception as e:
        logger.error(f"[Word生成] 保存失败: {str(e)}", exc_info=True)
        raise


def save_as_txt(content, output_path):
    """
    生成标准TXT文档
    改进点：保留原始段落结构

    :param content: 段落列表（保持原有换行符）
    :param output_path: 输出文件路径
    """
    logger.info(f"[TXT生成] 开始保存文件: {output_path}")
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            # 直接写入段落列表（保留所有换行符）
            f.writelines(content)
        logger.debug(f"[TXT生成] 保存成功 | 段落数: {len(content)}")
    except Exception as e:
        logger.error(f"[TXT生成] 保存失败: {str(e)}", exc_info=True)
        raise