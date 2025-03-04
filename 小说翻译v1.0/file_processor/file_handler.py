# file_processor\file_handler.py
"""
文件处理模块
功能：实现文档的智能分块、格式解析和输出保存
核心机制：
- 支持DOCX/TXT格式的解析与生成
- 动态分块算法适配多语言特性
- 自动检测文件编码格式
- 保留文档结构并生成符合排版规范的输出文件
"""

import os
import re
import hashlib
import logging
import uuid
import chardet
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 导入配置文件
from config.settings import FILE_HANDLER_CONFIG

# 初始化模块级日志记录器
logger = logging.getLogger("FileHandler")


def get_file_hash(file_path):
    """
    生成文件唯一特征码（MD5哈希值）

    实现原理：
    - 读取文件二进制内容计算MD5摘要
    - 用于会话缓存管理和文件版本标识

    :param file_path: 文件绝对路径
    :return: 32位十六进制哈希字符串
    :raises IOError: 文件读取失败时抛出
    """
    logger.debug(f"[特征码生成] 开始处理文件: {file_path}")
    try:
        with open(file_path, 'rb') as f:
            # 计算文件内容的MD5哈希值
            file_hash = hashlib.md5(f.read()).hexdigest()
            logger.info(f"[特征码生成] 成功生成 | 文件: {os.path.basename(file_path)} | 哈希: {file_hash[:8]}...")
            return file_hash
    except Exception as e:
        logger.error(f"[特征码生成] 处理失败: {str(e)}", exc_info=True)
        raise


def dynamic_split(text, target_lang, max_tokens=FILE_HANDLER_CONFIG["CHUNKING"]["DEFAULT_MAX_TOKENS"]):
    """
    增强型动态分块算法

    核心特性：
    - 语言自适应的分句规则
    - 连接词感知的分块保护
    - 亚洲语言缓冲系数调节
    - 上下文连贯性保持

    :param text: 原始文本内容
    :param target_lang: 目标语言代码（zh/en/ja/ko）
    :param max_tokens: 单块最大字节数（默认从配置文件中读取）
    :return: 分块后的文本列表
    """
    logger.info(f"[智能分块] 启动分块处理 | 目标语言: {target_lang} | 最大长度: {max_tokens}字节")

    # 从配置文件中获取语言特定配置
    lang_config = FILE_HANDLER_CONFIG["CHUNKING"]["LANG_CONFIG"]
    # 获取当前语言配置，默认使用英语配置
    config = lang_config.get(target_lang, lang_config['en'])
    logger.debug(f"[智能分块] 应用分块配置: {config}")

    # 构建正则表达式模式
    sentence_end = config['sentence_end'] + r'\s*'  # 匹配句尾符号及后续空白
    connectors = config['connectors']

    # 处理连接词后的错误分割问题
    for conn in connectors:
        # 在连接词+标点处插入临时标记防止错误分割
        text = re.sub(rf'({conn}[。．!?])', rf'\1##NO_SPLIT_{uuid.uuid4().hex}##', text)

    # 执行初步分句
    sentences = re.split(sentence_end, text)

    # 还原临时标记为原始内容
    sentences = [s.replace(f'##NO_SPLIT_{uuid.uuid4().hex}##', '') for s in sentences]

    chunks = []
    current_chunk = []
    current_length = 0
    # 从配置文件中获取亚洲语言缓冲系数
    factor = FILE_HANDLER_CONFIG["CHUNKING"]["ASIAN_BUFFER_FACTOR"] if target_lang in ['ja', 'zh'] else 1.0

    # 二次分块处理
    for sent in sentences:
        sent = sent.strip()
        if not sent:
            continue

        # 计算当前句子字节长度
        sent_length = len(sent.encode('utf-8'))

        # 分块条件判断
        if current_length + sent_length > max_tokens * factor:
            if current_chunk:
                chunks.append(''.join(current_chunk))
                current_chunk = []
                current_length = 0

        current_chunk.append(sent)
        current_length += sent_length

    # 处理剩余内容
    if current_chunk:
        chunks.append(''.join(current_chunk))
    # logger.debug(f"[智能分块] 分块完成 | 总块数: {len(chunks)}")
    for i, chunk in enumerate(chunks, 1):
        logger.info(f"分块 {i}: {len(chunk)} 字符 / {len(chunk.encode('utf-8'))} 字节")
    logger.info(f"[智能分块] 处理完成 | 总块数: {len(chunks)}")

    return chunks


def extract_text_from_docx(docx_path):
    """
    解析DOCX文档内容

    实现特性：
    - 保留段落结构
    - 过滤空段落
    - 维护原始换行符

    :param docx_path: DOCX文件路径
    :return: 拼接后的纯文本内容
    :raises ValueError: 文件解析失败时抛出
    """
    logger.info(f"[DOCX解析] 开始处理文件: {docx_path}")
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        logger.debug(f"[DOCX解析] 成功解析 | 段落数: {len(full_text)}")
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"[DOCX解析] 处理失败: {str(e)}", exc_info=True)
        raise


def extract_text_from_txt(txt_path):
    """
    解析TXT文档内容（自动检测编码）

    实现特性：
    - 使用chardet自动检测编码
    - 处理非常用编码格式
    - 容错空文件处理

    :param txt_path: TXT文件路径
    :return: 解码后的文本内容
    :raises UnicodeDecodeError: 编码检测失败时抛出
    """
    logger.info(f"[TXT解析] 开始处理文件: {txt_path}")
    try:
        with open(txt_path, 'rb') as f:
            raw_data = f.read()
            # 自动检测文件编码
            encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
            text = raw_data.decode(encoding)
            logger.debug(f"[TXT解析] 编码检测结果: {encoding}")
            return text
    except Exception as e:
        logger.error(f"[TXT解析] 处理失败: {str(e)}", exc_info=True)
        raise


def save_as_word(content, output_path):
    """
    生成符合排版规范的Word文档

    排版规范：
    - 首行缩进24磅
    - 两端对齐
    - 西文Times New Roman字体
    - 东亚字符楷体
    - 固定行距1倍

    :param content: 内容列表（按分块顺序）
    :param output_path: 输出文件路径
    :raises PermissionError: 文件写入权限不足时抛出
    """
    logger.info(f"[Word生成] 开始保存文件: {output_path}")
    try:
        doc = Document()
        for chunk in content:
            for para_text in chunk.split('\n'):
                if para_text.strip():
                    p = doc.add_paragraph()
                    # 设置段落格式
                    p_format = p.paragraph_format
                    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
                    p_format.line_spacing = FILE_HANDLER_CONFIG["DOCX_STYLE"]["LINE_SPACING"]  # 固定行距
                    p_format.first_line_indent = FILE_HANDLER_CONFIG["DOCX_STYLE"]["INDENT"]  # 首行缩进

                    # 设置字体样式
                    run = p.add_run(para_text)
                    run.font.name = FILE_HANDLER_CONFIG["DOCX_STYLE"]["FONTS"]["western"]  # 西文字体
                    run.font.size = FILE_HANDLER_CONFIG["DOCX_STYLE"]["FONT_SIZE"]
                    # 东亚字符字体设置
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FILE_HANDLER_CONFIG["DOCX_STYLE"]["FONTS"]["east_asian"])

        doc.save(output_path)
        logger.debug(f"[Word生成] 保存成功 | 文件大小: {os.path.getsize(output_path) // 1024}KB")
    except Exception as e:
        logger.error(f"[Word生成] 保存失败: {str(e)}", exc_info=True)
        raise


def save_as_txt(content, output_path):
    """
    生成标准TXT文档

    实现特性：
    - UTF-8编码保证兼容性
    - 保留分块结构
    - 自动处理换行符

    :param content: 内容列表（按分块顺序）
    :param output_path: 输出文件路径
    """
    logger.info(f"[TXT生成] 开始保存文件: {output_path}")
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            # 用换行符连接各分块内容
            f.write('\n'.join(content))
        logger.debug(f"[TXT生成] 保存成功 | 文件大小: {os.path.getsize(output_path) // 1024}KB")
    except Exception as e:
        logger.error(f"[TXT生成] 保存失败: {str(e)}", exc_info=True)
        raise