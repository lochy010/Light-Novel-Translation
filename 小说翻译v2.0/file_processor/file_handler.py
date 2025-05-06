# file_processor/file_handler.py
"""
该文件负责处理文档文件的读取、内容提取、智能分块以及将处理后的内容保存为新文件。
主要功能包括：
- 计算文件的 MD5 哈希值，用于缓存管理和文件识别。
- 从 DOCX 和 TXT 文件中提取文本内容，TXT 支持自动编码检测。
- 实现一个动态分块算法 (`dynamic_split`)，将长文本分割成适合翻译 API 处理的块，同时考虑语言特性和上下文连贯性。
- 将翻译后的文本块列表保存为格式化的 DOCX 文件或纯文本 TXT 文件 (UTF-8 编码)。
"""

# 导入 os 模块，用于路径操作（如获取文件名、目录名）、检查文件/目录存在性、获取文件大小等。
import os
# 导入 re 模块，用于正则表达式操作，主要在分块算法中进行句子分割和模式匹配。
import re
# 导入 hashlib 模块，用于计算文件的 MD5 哈希值。
import hashlib
# 导入 logging 模块，用于记录文件处理过程中的信息、警告和错误。
import logging
# 导入 uuid 模块，用于生成唯一的标识符，在分块算法中临时标记以防止错误分割。
import uuid
# 导入 chardet 库，用于自动检测 TXT 文件的字符编码。
import chardet
# 导入 python-docx 库中的 Document 类，用于读取和创建 DOCX 文件。
from docx import Document
# 导入 python-docx 库中的 Pt 类，用于指定 Word 文档中的磅值（如字号、缩进）。
from docx.shared import Pt
# 导入 python-docx 库中的 qn 函数，用于处理 Word 文档中的限定名称（Qualified Name），设置东亚字体时需要。
from docx.oxml.ns import qn
# 导入 python-docx 库中的 WD_PARAGRAPH_ALIGNMENT 枚举，用于设置段落对齐方式。
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 导入配置文件中的文件处理相关设置
# settings 模块包含了应用程序所需的各种配置常量，这里主要用到分块和 DOCX 样式的配置。
from config.settings import FILE_HANDLER_CONFIG

# --- 初始化模块级日志记录器 ---
# 获取一个名为 'file_handler' 的日志记录器实例。
# 日志的格式化将在主程序入口 (main.py) 中统一设置。
logger = logging.getLogger(__name__) # 使用当前模块名作为记录器名称

# --- 文件哈希生成函数 ---
def get_file_hash(file_path):
    """
    计算指定文件的 MD5 哈希值，生成文件的唯一特征码。

    实现原理：
    - 以二进制读取模式 (`'rb'`) 打开文件。
    - 读取文件的全部二进制内容。
    - 使用 `hashlib.md5()` 计算内容的 MD5 摘要。
    - 返回 32 位的小写十六进制哈希字符串。

    常用于：
    - 缓存系统中作为会话标识符，区分不同文件或同一文件的不同版本。
    - 文件版本控制或完整性校验。

    参数:
        file_path (str): 需要计算哈希值的文件完整路径。

    返回:
        str: 计算得到的 32 位十六进制 MD5 哈希字符串。

    异常:
        FileNotFoundError: 如果指定的文件路径不存在。
        IOError: 如果在读取文件时发生权限问题或其他 OS 级别的 IO 错误。
        Exception: 如果在计算哈希过程中发生其他未预期的错误。
    """
    # 获取文件名用于日志记录
    base_filename = os.path.basename(file_path)
    logger.info(f"[Hash Generation] 开始为文件 '{base_filename}' 生成 MD5 哈希值...")
    logger.debug(f"[Hash Generation] 文件完整路径: {file_path}")
    try:
        # 以二进制读取模式打开文件 ('rb')
        with open(file_path, 'rb') as f:
            # 读取文件全部内容
            file_content = f.read()
            # 记录读取的内容长度
            content_length = len(file_content)
            logger.debug(f"[Hash Generation] 成功读取文件内容，长度: {content_length} 字节。")
            # 计算内容的 MD5 哈希值并获取十六进制表示
            file_hash = hashlib.md5(file_content).hexdigest()
            # 记录成功生成哈希值的日志
            logger.info(f"[Hash Generation] 文件 '{base_filename}' MD5 哈希生成成功: {file_hash}")
            # 返回哈希值
            return file_hash
    except FileNotFoundError:
        # 文件不存在，记录错误并重新抛出异常
        logger.error(f"[Hash Generation] 文件未找到: {file_path}")
        raise
    except OSError as e:
        # 捕获操作系统级别的错误（如权限问题）
        logger.error(f"[Hash Generation] 读取文件时发生 OS 错误: {e} (文件: {file_path})", exc_info=True)
        # 包装成 IOError 抛出，提供更明确的错误信息
        raise IOError(f"读取文件失败: {e}") from e
    except Exception as e:
        # 捕获其他所有未预期的异常
        logger.error(f"[Hash Generation] 生成哈希时发生未预期错误: {e} (文件: {file_path})", exc_info=True)
        # 重新抛出原始异常
        raise
# --- 结束文件哈希生成函数 ---

# --- 动态分块函数 ---
def dynamic_split(text, target_lang, max_tokens=FILE_HANDLER_CONFIG["CHUNKING"]["DEFAULT_MAX_TOKENS"]):
    """
    将输入的文本（字符串或段落列表）分割成适合翻译 API 处理的文本块列表。
    该算法会尝试在句子或段落边界进行分割，并考虑最大字节数限制。

    核心特性:
    - 支持字符串或段落列表作为输入。
    - 根据目标语言 (`target_lang`) 从配置中选择合适的分句标点正则表达式。
    - 使用配置中的连接词列表（如果存在）来避免在这些词后错误地分割句子。
    - 对亚洲语言（中日韩）应用缓冲系数 (`ASIAN_BUFFER_FACTOR`)，调整最大字节数限制，以更好地估算 token 数量。
    - 如果单个段落超过调整后的最大字节数，会尝试在该段落内部按句子进行进一步分割 (`_split_oversized_para`)。
    - 将分割后的段落/句子组合成块，确保每个块的总字节数（估算）不超过限制。
    - 记录详细的分块过程日志。

    参数:
        text (str or list): 待分块的原始文本。可以是包含换行符的单个字符串，或一个字符串列表（每个元素代表一个段落）。
        target_lang (str): 目标翻译语言的代码（例如 'zh', 'en', 'ja', 'ko'），用于选择分句规则和缓冲系数。
        max_tokens (int, optional): 每个块的最大字节数限制（近似值）。默认为从 `FILE_HANDLER_CONFIG` 中读取。

    返回:
        list[str]: 分割后的文本块列表。每个元素是一个字符串，包含一个或多个段落/句子，段落之间用换行符 `\n` 分隔。

    异常:
        TypeError: 如果输入的 `text` 参数不是字符串或列表类型。
    """
    logger.info("[Dynamic Split] --- 开始动态分块处理 ---")
    logger.info(f"[Dynamic Split] 目标语言: {target_lang}, 最大分块(字节): {max_tokens}")

    # --- 输入类型处理 ---
    paragraphs = [] # 用于存储待处理的段落列表
    if isinstance(text, str):
        # 如果输入是字符串，按行分割成段落列表，保留换行符用于后续处理
        paragraphs = text.splitlines(keepends=True)
        logger.debug("[Dynamic Split] 输入为字符串，已按行分割成段落列表。")
    elif isinstance(text, list):
        # 如果输入已经是列表，直接使用
        paragraphs = text
        logger.debug(f"[Dynamic Split] 输入为列表，包含 {len(paragraphs)} 个段落。")
    else:
        # 如果输入类型不支持，记录错误并抛出 TypeError
        err_msg = f"不支持的输入类型: {type(text)}。应为字符串或列表。"
        logger.error(f"[Dynamic Split] {err_msg}")
        raise TypeError(err_msg)

    # --- 获取语言特定配置 ---
    # 从配置中获取分块相关的设置
    lang_config_map = FILE_HANDLER_CONFIG["CHUNKING"]["LANG_CONFIG"]
    # 设置默认配置（使用英文配置作为后备）
    default_lang_config = lang_config_map.get('en', {"sentence_end": r"[.!?…]", "connectors": []})
    # 获取目标语言的配置，如果找不到则使用默认配置
    lang_config = lang_config_map.get(target_lang, default_lang_config)
    # 如果使用了默认配置，记录警告
    if target_lang not in lang_config_map:
         logger.warning(f"[Dynamic Split] 未找到目标语言 '{target_lang}' 的特定配置，将使用默认配置。")
    # 记录使用的分句规则和连接词
    logger.info(f"[Dynamic Split] 应用分块配置: SentenceEnd='{lang_config['sentence_end']}', Connectors={lang_config.get('connectors', [])}")

    # --- 初始化变量 ---
    chunks = []                 # 存储最终分块结果的列表
    current_chunk_paras = []    # 当前正在构建的块所包含的段落/子段落列表
    current_chunk_bytes = 0     # 当前正在构建的块的累计字节数
    # 获取亚洲语言缓冲系数
    asian_buffer_factor = FILE_HANDLER_CONFIG["CHUNKING"]["ASIAN_BUFFER_FACTOR"]
    # 判断是否应用缓冲系数
    buffer_factor = asian_buffer_factor if target_lang in ['ja', 'zh', 'ko'] else 1.0
    # 计算调整后的最大字节数限制
    max_chunk_bytes_adjusted = max_tokens * buffer_factor
    logger.info(f"[Dynamic Split] 亚洲语言缓冲系数: {buffer_factor:.1f}, 调整后最大分块(字节): {max_chunk_bytes_adjusted:.0f}")

    # --- 遍历段落进行分块 ---
    para_index = 0 # 段落索引
    while para_index < len(paragraphs):
        # 获取当前段落并去除首尾空白（但不移除内部空白）
        para = paragraphs[para_index].strip()
        # 记录原始索引，便于日志追踪
        current_original_index = para_index
        # 移动到下一个段落索引
        para_index += 1

        # 跳过空段落
        if not para:
            logger.debug(f"[Dynamic Split] 跳过空段落 (原索引 {current_original_index})。")
            continue

        # 计算当前段落的 UTF-8 编码字节数
        para_bytes = len(para.encode('utf-8'))
        logger.debug(f"[Dynamic Split] 处理段落 (原索引 {current_original_index}, 字节: {para_bytes}) | 当前块字节: {current_chunk_bytes}")

        # --- 处理超大段落 ---
        # 如果当前段落字节数超过调整后的限制
        if para_bytes > max_chunk_bytes_adjusted:
            logger.warning(f"[Dynamic Split] 段落 (原索引 {current_original_index}, 字节: {para_bytes}) 超过最大分块限制 ({max_chunk_bytes_adjusted:.0f})，尝试句子级分割。")
            # 调用内部函数按句子分割超大段落
            sub_paras = _split_oversized_para(para, target_lang, max_tokens, lang_config, buffer_factor)
            logger.info(f"[Dynamic Split] 超大段落 (原索引 {current_original_index}) 被分割成 {len(sub_paras)} 个子段落。")

            # 遍历分割后的子段落
            for sub_index, sub in enumerate(sub_paras):
                # 去除子段落首尾空白
                sub = sub.strip()
                # 跳过空的子段落
                if not sub:
                    logger.debug(f"[Dynamic Split] 跳过分割后的空子段落 (原段落索引 {current_original_index}, 子段落 {sub_index+1})。")
                    continue
                # 计算子段落字节数
                sub_bytes = len(sub.encode('utf-8'))
                logger.debug(f"[Dynamic Split] 处理子段落 {current_original_index}.{sub_index+1} (字节: {sub_bytes}) | 当前块字节: {current_chunk_bytes}")

                # 检查将当前子段落加入块后是否会超长
                # （仅当块中已有内容时才需要检查，避免空块直接放入超长子段落的问题，虽然理论上_split_oversized_para会处理）
                if current_chunk_bytes > 0 and current_chunk_bytes + sub_bytes > max_chunk_bytes_adjusted:
                    # 如果超长，将当前块的内容合并成字符串，添加到 chunks 列表
                    chunk_content = '\n'.join(current_chunk_paras)
                    chunks.append(chunk_content)
                    logger.info(f"[Dynamic Split] 创建新块 (块 {len(chunks)}) - 原因: 添加子段落 {current_original_index}.{sub_index+1} 超长。")
                    logger.debug(f"[Dynamic Split]   块 {len(chunks)} 内容: {len(current_chunk_paras)} 段/子段, {current_chunk_bytes} 字节。")
                    # 重置当前块，并将当前子段落作为新块的第一个元素
                    current_chunk_paras = [sub]
                    current_chunk_bytes = sub_bytes
                else:
                    # 如果未超长，将子段落添加到当前块，并累加字节数
                    current_chunk_paras.append(sub)
                    current_chunk_bytes += sub_bytes
                    logger.debug(f"[Dynamic Split] 子段落 {current_original_index}.{sub_index+1} 已添加到当前块。当前块字节: {current_chunk_bytes}")
        # --- 处理普通段落 (未超长) ---
        else:
            # 检查将当前段落加入块后是否会超长
            # （仅当块中已有内容时才需要检查）
            if current_chunk_bytes > 0 and current_chunk_bytes + para_bytes > max_chunk_bytes_adjusted:
                # 如果超长，将当前块的内容合并成字符串，添加到 chunks 列表
                chunk_content = '\n'.join(current_chunk_paras)
                chunks.append(chunk_content)
                logger.info(f"[Dynamic Split] 创建新块 (块 {len(chunks)}) - 原因: 添加段落 (原索引 {current_original_index}) 超长。")
                logger.debug(f"[Dynamic Split]   块 {len(chunks)} 内容: {len(current_chunk_paras)} 段/子段, {current_chunk_bytes} 字节。")
                # 重置当前块，并将当前段落作为新块的第一个元素
                current_chunk_paras = [para]
                current_chunk_bytes = para_bytes
            else:
                # 如果未超长，将段落添加到当前块，并累加字节数
                current_chunk_paras.append(para)
                current_chunk_bytes += para_bytes
                logger.debug(f"[Dynamic Split] 段落 (原索引 {current_original_index}) 已添加到当前块。当前块字节: {current_chunk_bytes}")

    # --- 处理循环结束后剩余的块 ---
    # 如果循环结束后 current_chunk_paras 中还有内容，将其作为最后一个块添加
    if current_chunk_paras:
        chunk_content = '\n'.join(current_chunk_paras)
        chunks.append(chunk_content)
        logger.info(f"[Dynamic Split] 添加最后剩余的块 (块 {len(chunks)})。")
        logger.debug(f"[Dynamic Split]   块 {len(chunks)} 内容: {len(current_chunk_paras)} 段/子段, {current_chunk_bytes} 字节。")

    # --- 记录总结信息 ---
    total_paras_input = len(paragraphs) # 输入的原始段落/行数
    total_chunks_created = len(chunks)  # 生成的总块数
    logger.info(f"[Dynamic Split] --- 动态分块处理完成 ---")
    logger.info(f"[Dynamic Split] 原始输入段落/行数: {total_paras_input} -> 生成分块数: {total_chunks_created}")
    # 记录每个生成块的字节数（用于调试）
    for i, chunk_content in enumerate(chunks):
        chunk_bytes = len(chunk_content.encode('utf-8'))
        logger.debug(f"[Dynamic Split]   块 {i+1}/{total_chunks_created} | 字节: {chunk_bytes}")
    # 返回分块结果列表
    return chunks
# --- 结束动态分块函数 ---

# --- 超大段落分割辅助函数 ---
def _split_oversized_para(para, target_lang, max_tokens, config, buffer_factor):
    """
    内部辅助函数，用于将单个超大的段落按句子分割成多个子段落。
    该函数会尝试在句子末尾分割，并考虑连接词，然后将短句子合并，尽量使每个子段落接近但不超过最大字节数限制。

    参数:
        para (str): 需要分割的超大段落文本。
        target_lang (str): 目标语言代码，用于日志。
        max_tokens (int): 每个块的最大字节数限制（基准值）。
        config (dict): 包含该语言 'sentence_end' 正则表达式和 'connectors' 列表的配置字典。
        buffer_factor (float): 应用于 max_tokens 的缓冲系数。

    返回:
        list[str]: 分割和合并后的子段落列表。
    """
    log_prefix = "[Oversized Para Split]" # 日志前缀
    logger.info(f"{log_prefix} 开始分割超大段落...")
    # 记录部分原始内容用于调试，替换换行符以便单行显示
    logger.debug(f"{log_prefix} 原始内容 (前100字符): '{para[:100].replace(chr(10), '')}...'")

    # 获取句子结束符的正则表达式和连接词列表
    sentence_end_pattern = config['sentence_end'] + r'\s*' # 在结束符后匹配零个或多个空白符
    connectors = config.get('connectors', [])
    # 生成一个唯一的临时标记，用于防止在连接词后错误地分割
    no_split_marker = f"__NOSPLIT_{uuid.uuid4().hex[:8]}__"
    logger.debug(f"{log_prefix} 使用临时防分割标记: {no_split_marker}")

    # --- 防止在连接词后分割 ---
    temp_para = para # 操作副本
    markers_added = 0 # 记录添加的标记数量
    if connectors:
        logger.debug(f"{log_prefix} 检查连接词以防止错误分割...")
        for conn in connectors:
            # 构建正则表达式，匹配 "连接词 + 空白 + 句子结束符 + 空白"
            pattern = rf'({re.escape(conn)}\s*{sentence_end_pattern})'
            # 记录替换前的标记数量
            original_count = temp_para.count(no_split_marker)
            # 使用 re.sub 查找匹配项，并在匹配项后添加 no_split_marker
            # lambda m: m.group(1) + no_split_marker 保留原始匹配内容并追加标记
            temp_para = re.sub(pattern, lambda m: m.group(1) + no_split_marker, temp_para, flags=re.IGNORECASE)
            # 记录替换后的标记数量
            new_count = temp_para.count(no_split_marker)
            # 如果数量增加，说明添加了标记
            if new_count > original_count:
                markers_added += new_count - original_count
                logger.debug(f"{log_prefix}   - 为连接词 '{conn}' 添加了 {new_count - original_count} 个标记。")
        if markers_added > 0: logger.debug(f"{log_prefix} 共添加了 {markers_added} 个防分割标记。")
        else: logger.debug(f"{log_prefix} 未找到需要添加防分割标记的连接词。")
    else: logger.debug(f"{log_prefix} 未配置连接词，跳过防分割标记添加。")

    # --- 按句子分割 ---
    logger.debug(f"{log_prefix} 使用正则 '{sentence_end_pattern}' 进行句子分割...")
    # 使用 re.split 进行分割，括号 () 使分隔符本身也包含在结果列表中
    parts = re.split(f'({sentence_end_pattern})', temp_para)
    sentences = [] # 存储初步分割的句子
    i = 0
    # re.split 的结果是 [text, delim, text, delim, ...] 的形式，需要两两组合
    while i < len(parts):
        text_part = parts[i]                  # 文本部分
        punct_part = parts[i+1] if i+1 < len(parts) else "" # 分隔符部分（可能不存在）
        # 只有当文本或分隔符部分非空时才组合添加
        if text_part or punct_part: sentences.append(text_part + punct_part)
        i += 2 # 步长为 2
    logger.info(f"{log_prefix} 初步分割得到 {len(sentences)} 个句子/部分。")
    # 如果分割后没有句子，可能意味着段落结构特殊或正则有问题，返回原段落
    if not sentences:
        logger.warning(f"{log_prefix} 句子分割未能产生任何部分，将返回原段落。")
        return [para]

    # --- 恢复标记并清理空句子 ---
    restored_sentences = [] # 存储恢复标记并清理后的句子
    markers_removed = 0     # 记录移除的标记数量
    for s in sentences:
        original_s = s # 保存原始带标记的句子
        # 移除临时标记
        s_restored = s.replace(no_split_marker, '')
        # 如果移除标记后句子内容有变化，增加移除计数
        if s_restored != original_s: markers_removed += original_s.count(no_split_marker)
        # 去除句子首尾空白，如果句子非空，则添加到最终列表
        if s_restored.strip(): restored_sentences.append(s_restored.strip())
    if markers_removed > 0: logger.debug(f"{log_prefix} 成功移除了 {markers_removed} 个防分割标记。")
    logger.debug(f"{log_prefix} 清理空句子后剩余 {len(restored_sentences)} 个有效句子。")
    # 如果清理后没有有效句子，返回原段落
    if not restored_sentences:
        logger.warning(f"{log_prefix} 还原标记和清理后无有效句子，将返回原段落。")
        return [para]

    # --- 合并短句子 ---
    final_sub_paras = []         # 存储最终合并后的子段落列表
    current_sub_para_parts = []  # 当前正在构建的子段落包含的句子列表
    current_sub_para_bytes = 0   # 当前正在构建的子段落的累计字节数
    # 计算子段落的最大字节数限制
    max_bytes = max_tokens * buffer_factor

    logger.info(f"{log_prefix} 开始合并短句子，目标最大字节数: {max_bytes:.0f}")
    # 遍历清理后的有效句子
    for i, sentence in enumerate(restored_sentences):
        # 计算当前句子的字节数
        sentence_bytes = len(sentence.encode('utf-8'))
        logger.debug(f"{log_prefix} 处理句子 {i+1}/{len(restored_sentences)} (字节: {sentence_bytes}) | 当前子段字节: {current_sub_para_bytes}")

        # --- 处理单个句子超长的情况 ---
        # 理论上这不应该发生，除非句子本身就巨大无比
        if sentence_bytes > max_bytes:
             logger.warning(f"{log_prefix} 单个句子 (索引 {i+1}, 字节: {sentence_bytes}) 仍然超过最大长度限制 ({max_bytes:.0f})。将强制作为一个单独的子段落。")
             # 如果当前子段落有内容，先将其保存
             if current_sub_para_parts:
                 # 使用空格连接句子组成子段落
                 final_sub_paras.append(" ".join(current_sub_para_parts))
                 logger.debug(f"{log_prefix}   - 保存了累积的子段落 (字节: {current_sub_para_bytes})。")
                 # 重置当前子段落
                 current_sub_para_parts = []; current_sub_para_bytes = 0
             # 将这个超长的单句子直接作为一个子段落添加
             final_sub_paras.append(sentence)
             logger.debug(f"{log_prefix}   - 超长句子 {i+1} 已单独添加。")
             # 继续处理下一个句子
             continue

        # --- 检查添加当前句子是否会超长 ---
        # （仅当子段落已有内容时检查）
        if current_sub_para_bytes > 0 and current_sub_para_bytes + sentence_bytes > max_bytes:
            # 如果添加后会超长，将当前累积的子段落保存
            final_sub_paras.append(" ".join(current_sub_para_parts))
            logger.debug(f"{log_prefix} 创建新子段落 (原因: 添加句子 {i+1} 超长)。上一子段字节: {current_sub_para_bytes}")
            # 开始新的子段落，并将当前句子作为第一个元素
            current_sub_para_parts = [sentence]
            current_sub_para_bytes = sentence_bytes
        else:
            # 如果未超长，将当前句子添加到子段落，并累加字节数
            current_sub_para_parts.append(sentence)
            current_sub_para_bytes += sentence_bytes
            logger.debug(f"{log_prefix} 句子 {i+1} 已添加到当前子段落。当前子段字节: {current_sub_para_bytes}")

    # --- 处理循环结束后剩余的子段落 ---
    if current_sub_para_parts:
        final_sub_paras.append(" ".join(current_sub_para_parts))
        logger.debug(f"{log_prefix} 添加最后剩余的子段落 (字节: {current_sub_para_bytes})。")

    # --- 记录总结信息 ---
    logger.info(f"{log_prefix} 超大段落分割和合并完成，最终生成 {len(final_sub_paras)} 个子段落。")
    # 记录每个子段落的字节数（用于调试）
    for i, sub in enumerate(final_sub_paras):
         sub_bytes = len(sub.encode('utf-8'))
         logger.debug(f"{log_prefix}   子段落 {i+1}/{len(final_sub_paras)} | 字节: {sub_bytes}")
    # 返回子段落列表
    return final_sub_paras
# --- 结束超大段落分割辅助函数 ---

# --- DOCX 文本提取函数 ---
def extract_text_from_docx(docx_path):
    """
    从 DOCX 文件中提取所有段落的文本内容。

    参数:
        docx_path (str): DOCX 文件的完整路径。

    返回:
        list[str]: 包含文档中所有非空段落文本的列表。每个元素是一个段落的字符串。

    异常:
        FileNotFoundError: 如果指定的 DOCX 文件不存在。
        RuntimeError: 如果在解析 DOCX 文件过程中发生 python-docx 库相关的或其他未预期错误。
    """
    # 获取文件名用于日志
    base_filename = os.path.basename(docx_path)
    logger.info(f"[DOCX Extract] 开始解析 DOCX 文件: {base_filename}")
    logger.debug(f"[DOCX Extract] 文件路径: {docx_path}")
    paragraphs = [] # 存储提取到的段落文本
    try:
        # 使用 python-docx 打开文档
        doc = Document(docx_path)
        logger.debug(f"[DOCX Extract] Document 对象创建成功。开始遍历段落...")
        para_count = 0         # 总段落计数
        empty_para_count = 0   # 空段落计数
        # 遍历文档中的所有段落
        for i, para in enumerate(doc.paragraphs):
            para_count += 1
            # 获取段落的文本内容
            text = para.text
            # 检查段落文本去除空白后是否非空
            if text.strip():
                # 如果非空，添加到结果列表
                paragraphs.append(text)
                # 记录前几个段落的内容（用于调试）
                if len(paragraphs) <= 3: logger.debug(f"[DOCX Extract]   - 段落 {i+1}: '{text[:50].replace(chr(10), '')}...'")
            else:
                # 如果为空段落，增加空段落计数
                empty_para_count += 1
        # 记录解析完成的日志和统计信息
        logger.info(f"[DOCX Extract] DOCX 文件解析完成。总段落数: {para_count}, 有效段落数: {len(paragraphs)}, 空段落数: {empty_para_count}")
        # 返回包含有效段落文本的列表
        return paragraphs
    except FileNotFoundError:
        # 文件不存在，记录错误并重新抛出
        logger.error(f"[DOCX Extract] 文件未找到: {docx_path}")
        raise
    except Exception as e:
        # 捕获解析过程中的其他异常
        logger.error(f"[DOCX Extract] 解析 DOCX 文件时发生错误: {e} (文件: {base_filename})", exc_info=True)
        # 包装成 RuntimeError 抛出
        raise RuntimeError(f"解析 DOCX 文件失败: {e}") from e
# --- 结束 DOCX 文本提取函数 ---

# --- TXT 文本提取函数 ---
def extract_text_from_txt(txt_path):
    """
    从 TXT 文件中提取文本内容，并自动检测文件编码。

    执行流程：
    1.  以二进制模式 (`'rb'`) 读取文件全部内容。如果文件为空，返回空列表。
    2.  检查文件开头的 BOM (Byte Order Mark) 来确定编码 (UTF-32 LE/BE, UTF-16 LE/BE, UTF-8 SIG)。
    3.  如果未检测到 BOM，使用 `chardet` 库尝试检测编码。
    4.  构建一个包含多种常用编码的优先级列表（BOM 检测结果、UTF-8、chardet 检测结果、GB18030、GBK、Big5 等）。
    5.  按优先级顺序尝试使用列表中的编码解码文件内容。
    6.  一旦解码成功，记录使用的编码，并停止尝试。
    7.  如果所有优先级编码都失败，尝试使用 UTF-8 编码并忽略错误 (`errors='ignore'`) 进行最后解码，并记录警告。
    8.  如果最终解码也失败，抛出 `UnicodeDecodeError`。
    9.  将成功解码的文本内容按行分割（保留换行符），并返回行列表。

    参数:
        txt_path (str): TXT 文件的完整路径。

    返回:
        list[str]: 包含文件中所有行的列表。每个元素是一行文本（可能包含末尾的换行符）。

    异常:
        FileNotFoundError: 如果指定的 TXT 文件不存在。
        UnicodeDecodeError: 如果尝试了所有已知编码都无法成功解码文件。
        RuntimeError: 如果在读取或解析过程中发生其他未预期错误。
    """
    # 获取文件名用于日志
    base_filename = os.path.basename(txt_path)
    logger.info(f"[TXT Extract] 开始解析 TXT 文件: {base_filename}")
    logger.debug(f"[TXT Extract] 文件路径: {txt_path}")

    try:
        # --- 读取二进制内容 ---
        logger.debug("[TXT Extract] 以二进制模式读取文件...")
        with open(txt_path, 'rb') as f: raw_data = f.read()
        # 检查文件是否为空
        if not raw_data:
            logger.warning(f"[TXT Extract] 文件 '{base_filename}' 为空，返回空列表。")
            return []
        logger.info(f"[TXT Extract] 文件读取成功，大小: {len(raw_data)} 字节。")

        # --- 检测 BOM ---
        logger.debug("[TXT Extract] 检测文件 BOM...")
        # 定义常见的 BOM 及其对应的编码名称
        bom_encodings = [
            (b'\xff\xfe\x00\x00', 'utf-32-le'), (b'\x00\x00\xfe\xff', 'utf-32-be'),
            (b'\xff\xfe', 'utf-16-le'), (b'\xfe\xff', 'utf-16-be'),
            (b'\xef\xbb\xbf', 'utf-8-sig'), # UTF-8 with BOM
        ]
        detected_encoding_bom = None # 存储通过 BOM 检测到的编码
        bom_length = 0              # 存储 BOM 的长度
        # 遍历 BOM 列表
        for bom, encoding in bom_encodings:
            # 检查文件内容是否以当前 BOM 开头
            if raw_data.startswith(bom):
                detected_encoding_bom = encoding # 记录检测到的编码
                bom_length = len(bom)          # 记录 BOM 长度
                logger.info(f"[TXT Extract] 检测到 BOM，编码确定为: {encoding}")
                # 从原始数据中移除 BOM 部分
                raw_data = raw_data[bom_length:]
                logger.debug(f"[TXT Extract] 已从原始数据中移除 BOM (长度: {bom_length})。")
                break # 找到 BOM 后即可停止检测
        if not detected_encoding_bom: logger.debug("[TXT Extract] 未检测到 BOM。")

        # --- 使用 chardet 检测 (如果未检测到 BOM) ---
        detected_encoding_chardet = None # 存储 chardet 检测结果
        used_chardet = False            # 标记是否使用了 chardet
        if not detected_encoding_bom:
            used_chardet = True
            logger.info("[TXT Extract] 未检测到BOM，尝试使用 chardet 进行编码检测...")
            # 调用 chardet.detect
            detection = chardet.detect(raw_data)
            detected_encoding_chardet = detection['encoding'] # 获取检测到的编码名称
            confidence = detection['confidence']              # 获取置信度
            logger.info(f"[TXT Extract] chardet 检测结果: {detected_encoding_chardet} (置信度: {confidence:.2f})")
            # 如果置信度较低，记录警告
            if confidence < 0.7: logger.warning(f"[TXT Extract] chardet 检测编码置信度较低 ({confidence:.2f})，结果可能不准确。")

        # --- 构建编码尝试列表 ---
        logger.debug("[TXT Extract] 构建编码尝试列表...")
        # 定义编码的优先级顺序
        encoding_priority_list = [
            detected_encoding_bom,      # 1. BOM 检测结果 (最高优先级)
            'utf-8',                    # 2. UTF-8 (常用)
            detected_encoding_chardet,  # 3. Chardet 检测结果
            'utf-8-sig',                # 4. UTF-8 with BOM (以防万一)
            'gb18030',                  # 5. 简体中文
            'gbk',                      # 6. 简体中文扩展
            'big5',                     # 7. 繁体中文
            'cp932', 'shift_jis', 'euc-jp', 'iso-2022-jp', # 8. 日语常用编码
            'euc-kr', 'cp949',          # 9. 韩语常用编码
            'cp1252', 'latin-1',        # 10. 西欧语言常用编码
        ]
        # 去重并过滤掉 None 值，生成最终的尝试列表
        encodings_to_try = []
        seen_encodings = set()
        for enc in encoding_priority_list:
            if enc and isinstance(enc, str): # 确保是有效的字符串编码名称
                enc_lower = enc.lower() # 转小写进行去重判断
                if enc_lower not in seen_encodings:
                    encodings_to_try.append(enc)
                    seen_encodings.add(enc_lower)
        logger.info(f"[TXT Extract] 将按以下顺序尝试解码: {encodings_to_try}")

        # --- 尝试解码 ---
        decoded_text = None # 存储解码后的文本
        used_encoding = None # 存储最终成功使用的编码
        # 遍历编码尝试列表
        for enc in encodings_to_try:
            try:
                logger.debug(f"[TXT Extract] 尝试使用编码 '{enc}' 解码...")
                # 尝试使用当前编码解码，使用 'strict' 错误处理（遇到错误即失败）
                decoded_text = raw_data.decode(enc, errors='strict')
                used_encoding = enc # 记录成功的编码
                logger.info(f"[TXT Extract] 文件成功使用编码 '{used_encoding}' 解码。")
                # 如果使用了 chardet 但最终编码与 chardet 首选不同，记录提示
                if used_chardet and detected_encoding_chardet and used_encoding.lower() != detected_encoding_chardet.lower():
                    logger.info(f"[TXT Extract] (注意: 最终使用编码 '{used_encoding}' 与 chardet 首选 '{detected_encoding_chardet}' 不同)")
                break # 解码成功，跳出循环
            except (UnicodeDecodeError, LookupError) as decode_error:
                # 如果解码失败（编码错误或不支持的编码），记录调试信息，继续尝试下一个
                logger.debug(f"[TXT Extract] 使用编码 '{enc}' 解码失败: {decode_error}")
                continue

        # --- 处理解码失败或最后尝试 ---
        # 如果遍历完所有优先级编码都未能成功解码
        if decoded_text is None:
            logger.error(f"[TXT Extract] 尝试了所有优先级编码，均无法解码文件 '{base_filename}'。")
            try:
                 # 最后尝试使用 UTF-8 并忽略错误
                 logger.warning("[TXT Extract] 尝试使用 'utf-8' 配合 'ignore' 错误处理进行最后解码...")
                 decoded_text = raw_data.decode('utf-8', errors='ignore')
                 used_encoding = 'utf-8 (errors=ignore)' # 标记使用了 ignore 模式
                 logger.warning(f"[TXT Extract] 使用 '{used_encoding}' 解码完成，但可能存在数据丢失。")
            except Exception as final_decode_error:
                 # 如果连最后尝试都失败，记录关键错误并抛出异常
                 logger.critical(f"[TXT Extract] 最终解码尝试也失败: {final_decode_error}", exc_info=True)
                 raise UnicodeDecodeError(f"无法使用任何已知编码解码文件 {txt_path}。请检查文件编码或内容。") from final_decode_error

        # --- 分割文本并返回 ---
        logger.debug("[TXT Extract] 解码完成，按行分割文本...")
        # 使用 splitlines(keepends=True) 按行分割，保留行尾的换行符
        # 保留换行符有助于 dynamic_split 更准确地处理原始格式
        paragraphs = decoded_text.splitlines(keepends=True)
        original_line_count = len(paragraphs)
        # 注意：这里不过滤空行，将空行处理交给 dynamic_split
        logger.info(f"[TXT Extract] 文本已按行分割。总行数: {original_line_count}")
        logger.info(f"[TXT Extract] --- TXT 文件解析完成 ---")
        # 返回包含所有行的列表
        return paragraphs

    except FileNotFoundError:
        # 文件不存在，记录错误并重新抛出
        logger.error(f"[TXT Extract] 文件未找到: {txt_path}")
        raise
    except Exception as e:
        # 捕获解析过程中的其他未知错误
        logger.error(f"[TXT Extract] 解析 TXT 文件时发生未知错误: {e} (文件: {base_filename})", exc_info=True)
        # 包装成 RuntimeError 抛出
        raise RuntimeError(f"解析 TXT 文件失败: {e}") from e
# --- 结束 TXT 文本提取函数 ---

# --- 保存为 Word 函数 ---
def save_as_word(content_list, output_path):
    """
    将翻译后的文本内容列表生成为一个格式化的 Word 文档 (.docx)。

    执行流程：
    1.  创建一个新的 Word 文档对象 (`Document`)。
    2.  从 `FILE_HANDLER_CONFIG` 获取文档样式配置（字体、字号、行距、缩进等）。
    3.  记录应用的排版规范日志。
    4.  遍历输入的 `content_list`（假设每个元素是一个翻译后的块，块内可能包含由 `\n` 分隔的段落）。
    5.  对于每个块，按换行符 `\n` 分割成段落。
    6.  对于每个非空段落：
        -   在 Word 文档中添加一个新段落 (`doc.add_paragraph()`)。
        -   获取该段落的格式对象 (`paragraph_format`)。
        -   设置段落对齐方式（两端对齐）、行距和首行缩进。
        -   在该段落中添加一个文本运行 (`run = p.add_run(text)`)。
        -   获取文本运行的字体对象 (`font`)。
        -   设置字体大小和西文字体 (`font.name`)。
        -   使用 `rPr.rFonts.set(qn('w:eastAsia'), ...)` 设置东亚字体。
    7.  记录添加到文档的总段落数。
    8.  调用 `doc.save(output_path)` 将文档保存到指定路径。
    9.  尝试获取保存后文件的大小并记录。

    参数:
        content_list (list[str]): 包含翻译后文本块的列表。每个字符串可能包含多个由 `\n` 分隔的段落。
        output_path (str): 输出的 Word 文档 (.docx) 的完整路径。

    异常:
        PermissionError: 如果没有写入指定输出路径的权限。
        RuntimeError: 如果在创建或保存 Word 文档过程中发生 python-docx 库相关的或其他未预期错误。
    """
    # 获取输出文件名用于日志
    base_filename = os.path.basename(output_path)
    logger.info(f"[Word Save] 开始将内容保存为 Word 文档: {base_filename}")
    logger.debug(f"[Word Save] 输出路径: {output_path}")
    logger.debug(f"[Word Save] 输入内容包含 {len(content_list)} 个块。")

    try:
        # 创建一个新的 Word 文档对象
        doc = Document()
        # --- 获取样式配置 ---
        style_config = FILE_HANDLER_CONFIG["DOCX_STYLE"]
        font_size = style_config.get("FONT_SIZE", Pt(12)) # 默认 12 磅
        font_western = style_config["FONTS"].get("western", "Calibri") # 默认 Calibri
        font_east_asian = style_config["FONTS"].get("east_asian", "Microsoft YaHei") # 默认微软雅黑
        line_spacing = style_config.get("LINE_SPACING", 1.0) # 默认单倍行距
        first_line_indent = style_config.get("INDENT", Pt(24)) # 默认首行缩进 24 磅
        paragraph_alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # 默认两端对齐

        # --- 记录应用的样式 ---
        logger.info("[Word Save] 应用排版规范:")
        logger.info(f"[Word Save]   - 字体: 西文 '{font_western}', 东亚 '{font_east_asian}'")
        logger.info(f"[Word Save]   - 字号: {font_size.pt:.1f} pt")
        logger.info(f"[Word Save]   - 行距: {line_spacing}")
        logger.info(f"[Word Save]   - 首行缩进: {first_line_indent.pt:.1f} pt")
        logger.info(f"[Word Save]   - 对齐方式: 两端对齐")

        # --- 遍历内容并添加到文档 ---
        total_paras_added = 0 # 记录添加到文档的有效段落数
        # 遍历输入的块列表
        for block_index, chunk in enumerate(content_list):
            # 将每个块按换行符分割成段落
            paragraphs_in_chunk = chunk.split('\n')
            logger.debug(f"[Word Save] 处理块 {block_index + 1}/{len(content_list)}，包含 {len(paragraphs_in_chunk)} 个段落。")
            # 遍历块内的段落
            for para_text in paragraphs_in_chunk:
                # 去除段落首尾空白
                para_text_stripped = para_text.strip()
                # 只处理非空段落
                if para_text_stripped:
                    total_paras_added += 1
                    # 在文档中添加新段落
                    p = doc.add_paragraph()
                    # 获取段落格式对象
                    p_format = p.paragraph_format
                    # 应用段落格式
                    p_format.alignment = paragraph_alignment # 对齐方式
                    p_format.line_spacing = line_spacing     # 行距
                    p_format.first_line_indent = first_line_indent # 首行缩进
                    # 在段落中添加文本运行
                    run = p.add_run(para_text_stripped)
                    # 获取字体对象
                    font = run.font
                    # 设置字体大小
                    font.size = font_size
                    # 设置西文字体（默认）
                    font.name = font_western
                    # 设置东亚字体（需要使用 qn）
                    # run._element.rPr 是获取运行属性的底层 OXML 元素
                    # .rFonts 是字体设置元素
                    # .set(qn('w:eastAsia'), font_east_asian) 设置东亚字体属性
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_east_asian)

        logger.info(f"[Word Save] 内容处理完成，共添加 {total_paras_added} 个有效段落到文档中。")
        logger.info("[Word Save] 正在执行文档保存操作...")
        # 将文档保存到指定的输出路径
        doc.save(output_path)
        # 尝试获取保存后的文件大小
        try:
            final_size = os.path.getsize(output_path)
            logger.info(f"[Word Save] Word 文档保存成功 | 大小: {final_size} 字节")
        except OSError:
            # 获取大小失败不影响结果，记录警告
            logger.warning("[Word Save] Word 文档已保存，但无法获取文件大小。")

    except PermissionError as e:
        # 捕获权限错误
        logger.error(f"[Word Save] 保存 Word 文档失败: 权限不足。路径: {output_path}", exc_info=True)
        # 重新抛出原始异常
        raise
    except Exception as e:
        # 捕获其他保存过程中的错误
        logger.error(f"[Word Save] 保存 Word 文档时发生未知错误: {e} (路径: {output_path})", exc_info=True)
        # 包装成 RuntimeError 抛出
        raise RuntimeError(f"保存 Word 文档失败: {e}") from e
# --- 结束保存为 Word 函数 ---

# --- 保存为 TXT 函数 ---
def save_as_txt(content_list, output_path):
    """
    将翻译后的文本内容列表生成为一个 UTF-8 编码的纯文本文件 (.txt)。

    执行流程：
    1.  使用换行符 (`\n`) 将 `content_list` 中的所有块连接成一个单一的字符串。块内原有的换行符会被保留。
    2.  以写入模式 (`'w'`) 和 UTF-8 编码打开指定的 `output_path` 文件。
    3.  将连接后的完整文本写入文件。
    4.  记录写入成功的日志和写入的字符数。
    5.  尝试获取保存后文件的大小并记录。

    参数:
        content_list (list[str]): 包含翻译后文本块的列表。每个字符串可能包含由 `\n` 分隔的段落。
        output_path (str): 输出的 TXT 文件 (.txt) 的完整路径。

    异常:
        PermissionError: 如果没有写入指定输出路径的权限。
        RuntimeError: 如果在写入文件过程中发生 IO 错误或其他未预期错误。
    """
    # 获取输出文件名用于日志
    base_filename = os.path.basename(output_path)
    logger.info(f"[TXT Save] 开始将内容保存为 TXT 文件: {base_filename}")
    logger.debug(f"[TXT Save] 输出路径: {output_path}")
    logger.debug(f"[TXT Save] 输入内容包含 {len(content_list)} 个块。")

    try:
        # 以写入模式 ('w') 和 UTF-8 编码打开文件
        # 'w' 模式会覆盖已存在的文件内容
        with open(output_path, 'w', encoding='utf-8') as f:
            # 使用换行符连接列表中的所有字符串（块）
            # 块内部原有的换行符会保留下来
            full_text = '\n'.join(content_list)
            # 将完整的文本写入文件，f.write() 返回写入的字符数
            chars_written = f.write(full_text)
            logger.info(f"[TXT Save] 内容已成功写入文件。写入字符数: {chars_written}")

        # 尝试获取保存后文件的大小
        try:
            final_size = os.path.getsize(output_path)
            logger.info(f"[TXT Save] TXT 文件保存成功 | 大小: {final_size} 字节")
        except OSError:
            # 获取大小失败记录警告
            logger.warning("[TXT Save] TXT 文件已保存，但无法获取文件大小。")

    except PermissionError as e:
        # 捕获权限错误
        logger.error(f"[TXT Save] 保存 TXT 文件失败: 权限不足。路径: {output_path}", exc_info=True)
        # 重新抛出原始异常
        raise
    except Exception as e:
        # 捕获其他写入过程中的错误
        logger.error(f"[TXT Save] 保存 TXT 文件时发生未知错误: {e} (路径: {output_path})", exc_info=True)
        # 包装成 RuntimeError 抛出
        raise RuntimeError(f"保存 TXT 文件失败: {e}") from e
# --- 结束保存为 TXT 函数 ---